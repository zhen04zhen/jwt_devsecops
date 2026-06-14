from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Helpers ──
def heading(text, level=1, color=RGBColor(0x1F, 0x49, 0x7D)):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
    return p

def body(text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def placeholder(label="[ 請插入截圖 ]"):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    run.font.italic = True
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row[c_idx].text = cell_text
            for run in row[c_idx].paragraphs[0].runs:
                run.font.size = Pt(10)
    return table

# ════════════════════════════════════════════
#  Title
# ════════════════════════════════════════════
title = doc.add_heading("DevSecOps Pipeline 作業報告", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph("作業 2 ── JWT 登入驗證 API 安全流水線")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in sub.runs:
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

repo_p = doc.add_paragraph()
repo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
repo_run = repo_p.add_run("GitHub Repo：https://github.com/zhen04zhen/jwt-devsecops")
repo_run.font.size = Pt(11)
repo_run.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)

doc.add_paragraph()

# ════════════════════════════════════════════
#  Section 1：專案說明
# ════════════════════════════════════════════
heading("一、專案說明", 1)

body("本專案為一個簡易的 JWT（JSON Web Token）登入驗證 REST API，"
     "使用 Python 與 Flask 框架開發，並整合 GitHub Actions 建立完整的 DevSecOps Pipeline，"
     "在每次程式碼提交時自動執行功能測試與多項安全掃描。")

doc.add_paragraph()
heading("1.1 使用語言與技術棧", 2)

add_table(
    ["元件", "技術 / 版本"],
    [
        ["程式語言",     "Python 3.12"],
        ["Web 框架",     "Flask 3.1.3"],
        ["JWT 套件",     "PyJWT 2.13.0"],
        ["密碼雜湊",     "bcrypt 4.1.3"],
        ["測試框架",     "pytest 9.0.3 + pytest-flask 1.3.0"],
        ["容器化",       "Docker（python:3.12-slim）"],
        ["CI/CD",        "GitHub Actions"],
    ]
)

doc.add_paragraph()
heading("1.2 API 端點", 2)

add_table(
    ["Method", "Endpoint", "說明"],
    [
        ["GET",  "/health",    "健康檢查"],
        ["POST", "/login",     "使用者登入，成功回傳 JWT Token"],
        ["GET",  "/protected", "受保護路由，需帶 Bearer Token 存取"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════
#  Section 2：Pipeline 說明
# ════════════════════════════════════════════
heading("二、Pipeline 說明", 1)

body("本專案的 GitHub Actions Pipeline 在每次 push 至 main 分支時自動觸發，"
     "共包含 11 個 Step，涵蓋環境建置、功能測試、安全掃描、容器打包與報告上傳。")

doc.add_paragraph()
heading("2.1 Pipeline 架構圖", 2)

diagram_path = r"C:\Users\zhen9\Desktop\DevSecOps\jwt-devsecops\pipeline_diagram.png"
if os.path.exists(diagram_path):
    doc.add_picture(diagram_path, width=Inches(6.2))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
heading("2.2 各工具職責說明", 2)

add_table(
    ["Step", "工具", "類型", "說明"],
    [
        ["1-3",  "actions/checkout、setup-python、pip", "環境建置", "取得程式碼、設定 Python 3.12、安裝依賴套件"],
        ["4",    "pytest",                              "功能測試", "執行 10 個測試案例，驗證 API 行為正確性"],
        ["5",    "Gitleaks",                            "機密掃描", "掃描 git commit history，偵測意外寫入程式碼的 secret"],
        ["6",    "ShiftLeft Scan",                      "SAST",     "靜態程式碼安全分析，偵測 SQL Injection、硬編碼密碼等問題"],
        ["7",    "pip-audit",                           "依賴掃描", "掃描 Python 套件是否有已知 CVE 漏洞"],
        ["8",    "docker build",                        "容器打包", "將應用程式打包成 Docker Image"],
        ["9",    "Trivy",                               "容器掃描", "掃描 Docker Image 中 OS 層與套件層的已知漏洞"],
        ["10",   "Dockle",                              "容器掃描", "檢查 Dockerfile 是否符合安全最佳實踐"],
        ["11",   "actions/upload-artifact",             "報告上傳", "上傳 pip-audit、Trivy、Dockle 掃描報告至 GitHub Actions Artifacts"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════
#  Section 3：測試過程
# ════════════════════════════════════════════
heading("三、Pipeline 測試過程", 1)

body("以下展示三次 Pipeline 執行紀錄，示範 Pipeline 如何及早發現問題。")

doc.add_paragraph()
heading("3.1 第一次執行：成功", 2)
body("程式碼狀態：乾淨，所有安全規範均遵守，套件版本皆為最新安全版本。")
body("預期結果：全部 11 個 Step 通過，Pipeline 顯示綠燈。")
doc.add_paragraph()
placeholder("[ 請插入第一次 Pipeline 成功的 Actions 截圖 ]")

doc.add_paragraph()
heading("3.2 第二次執行：故意失敗", 2)
body("改動內容：將 tests/test_protected.py 中的 test_access_without_token "
     "預期狀態碼從 401 改為 200。")
body("說明：不帶 Token 存取 /protected 端點，正確行為應回傳 401（未授權），"
     "但測試被改為預期 200（成功），造成測試失敗。")
body("Pipeline 行為：Step 4（Run tests）失敗後立即停止，後續安全掃描不再執行。")
doc.add_paragraph()
placeholder("[ 請插入第二次 Pipeline 失敗的 Actions 截圖（顯示 Run tests 紅燈）]")

doc.add_paragraph()
heading("3.3 第三次執行：修正後再次成功", 2)
body("改動內容：將 test_access_without_token 的預期狀態碼改回正確的 401。")
body("預期結果：全部 Step 恢復通過，Pipeline 再次顯示綠燈。")
doc.add_paragraph()
placeholder("[ 請插入第三次 Pipeline 成功的 Actions 截圖 ]")

doc.add_page_break()

# ════════════════════════════════════════════
#  Section 4：套件弱點修補
# ════════════════════════════════════════════
heading("四、套件弱點修補", 1)

body("在 Pipeline 執行過程中，pip-audit 掃描到以下套件存在已知漏洞，並進行修補。")

doc.add_paragraph()
heading("4.1 修補前（pip-audit 掃描結果）", 2)

add_table(
    ["套件", "原始版本", "CVE / PYSEC ID", "修補版本"],
    [
        ["pytest",  "8.2.2", "CVE-2025-71176",  "9.0.3"],
        ["PyJWT",   "2.8.0", "PYSEC-2026-120",  "2.13.0"],
        ["PyJWT",   "2.8.0", "PYSEC-2025-183",  "2.13.0"],
        ["PyJWT",   "2.8.0", "PYSEC-2026-175",  "2.13.0"],
        ["PyJWT",   "2.8.0", "PYSEC-2026-177",  "2.13.0"],
        ["PyJWT",   "2.8.0", "PYSEC-2026-178",  "2.13.0"],
        ["PyJWT",   "2.8.0", "PYSEC-2026-179",  "2.13.0"],
    ]
)

doc.add_paragraph()
placeholder("[ 請插入 pip-audit 顯示漏洞的 Actions 截圖 ]")

doc.add_paragraph()
heading("4.2 修補後（升級套件版本）", 2)
body("將 requirements.txt 中的版本升級如下：")

add_table(
    ["套件", "升級前", "升級後"],
    [
        ["pytest", "8.2.2",  "9.0.3"],
        ["PyJWT",  "2.8.0",  "2.13.0"],
        ["flask",  "3.0.3",  "3.1.3（配合 pytest-flask 依賴）"],
    ]
)

doc.add_paragraph()
body("升級後 pip-audit 掃描結果為零漏洞，Pipeline 恢復正常。")
doc.add_paragraph()
placeholder("[ 請插入 pip-audit 修補後無漏洞的 Actions 截圖 ]")

doc.add_page_break()

# ════════════════════════════════════════════
#  Section 5：掃描報告說明
# ════════════════════════════════════════════
heading("五、掃描報告說明", 1)

body("每次 Pipeline 執行完畢，Step 11 會將以下掃描報告上傳至 GitHub Actions Artifacts，"
     "可於 Actions 頁面下載查閱。")

doc.add_paragraph()
add_table(
    ["報告檔案", "產出工具", "內容說明"],
    [
        ["pip-audit-report.json", "pip-audit",  "列出所有 Python 依賴套件的 CVE 掃描結果，包含套件名稱、版本、漏洞 ID 與修補版本"],
        ["trivy-report.json",     "Trivy",       "列出 Docker Image 中 OS 套件與應用套件的漏洞，包含 CVE ID、嚴重程度（CRITICAL / HIGH / MEDIUM）與修補建議"],
        ["dockle-report.json",    "Dockle",      "列出 Dockerfile 不符合 CIS Docker Benchmark 的項目，例如是否以非 root 身份執行、是否設定 HEALTHCHECK 等"],
    ]
)

doc.add_paragraph()
heading("5.1 Artifacts 下載位置", 2)
body("GitHub Actions → 選擇任一 workflow run → 頁面底部 Artifacts 區塊 → 下載 security-reports.zip")
doc.add_paragraph()
placeholder("[ 請插入 GitHub Actions Artifacts 區塊截圖 ]")

doc.add_paragraph()
heading("5.2 Trivy 掃描結果說明", 2)
body("Trivy 掃描 Docker Image 時會同時掃描：\n"
     "• OS 層（Alpine / Debian 系統套件）\n"
     "• 應用層（Python 套件，與 pip-audit 互補）\n"
     "本專案設定 exit-code: 0，即使發現漏洞也不中斷 Pipeline，結果以報告形式保存。")

doc.add_paragraph()
heading("5.3 Dockle 掃描結果說明", 2)
body("Dockle 依據 CIS Docker Benchmark 檢查 Dockerfile 安全設定，本專案已針對以下項目進行修正：\n"
     "• 使用非 root 使用者（appuser）執行應用程式\n"
     "• 設定 HEALTHCHECK 指令\n"
     "• 使用 python:3.12-slim 固定版本 base image")

doc.add_paragraph()
placeholder("[ 請插入 trivy-report.json 或 dockle-report.json 內容截圖 ]")

# ════════════════════════════════════════════
#  Save
# ════════════════════════════════════════════
output_path = r"C:\Users\zhen9\Desktop\DevSecOps\DevSecOps_作業2報告.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
